import glob
import os
from tkinter import font # Seleciona fonte
from docx import Document # Usado para criar e salvar docx
from docx.shared import Inches # Seleciona o tamanho das imagens em polegadas
from docx.enum.text import WD_ALIGN_PARAGRAPH # Alinhar textos
from docx.shared import Pt # Seleciona o tamanho da fonte

#################### Funções ####################

def addTitle():
    # Coloca titulo no relatório
    p = document.add_paragraph('''Relatório fotográfico dos pontos, referente sistema de informação TV Prefeitura
    Nota fiscal número: XXX  emissão XXX Empenho nº XXX
    Contrato nº XXX – Processo licitatório nº XXX
    ''')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.style = document.styles['Normal']

def generateTable(imgsLeft):
    # Verifica a quantia restante da imagem
    imgsToPost = numImgs - imgsLeft

    numRows = 0 # Quantidade de linhas
    numCols = 0 # Quantidade de colunas
    maxRows = 4 # Quantidade de linhas máximas por pagina
    maxCols = 4 # Quantidade de colunas máximas por pagina

    # Filtro para saber quantas linhas e colunas desenhar, baseado na quantidade de imagens restantes
    if(imgsToPost >= maxRows+maxCols):
        imgsToPost = maxRows+maxCols
        numRows = maxRows
        numCols = maxCols
    elif(imgsToPost > maxCols):
        numCols = maxCols
        numRows = (imgsToPost // maxCols) * 2
    else:
        numCols = imgsToPost
        numRows = 2

    # Cria a tabela com base nos resultados do filtro
    table = document.add_table(rows=numRows, cols=numCols) 
    
    imgsPosted = 0 # Contador para saber quantas noticias faltam serem postadas na pagina
    currentRow = 0 # Marcador para saber qual linha esta sendo postado

    # Posta as imagens e o nome delas
    while(imgsPosted < imgsToPost): 
        img_cells = table.rows[currentRow].cells
        name_cells = table.rows[currentRow+1].cells
        for cols in range(numCols):
            # Adiciona imagem
            p1 = img_cells[cols].add_paragraph()
            r1 = p1.add_run()
            r1.add_picture("imagens/"+files[imgsPosted+imgsLeft], width=Inches(1.2))
            print("A imagem "+files[imgsPosted+imgsLeft]+" foi adicionada.")

            # Adiciona nome
            name_cells[cols].text = os.path.splitext(files[imgsPosted+imgsLeft])[0]

            imgsPosted += 1
        currentRow += 2
        numCols = imgsToPost - imgsPosted
    global imgCounter
    imgCounter += maxRows + maxCols    


#################################################

# Cria um documento em DOCX na memoria
document = Document('Template de arquivo/template.docx')

# Seleciona fonte e seu tamanho
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

print('Iniciando\n')

# Adiciona as imagens na lista files
files = []
images = glob.iglob('imagens/*.jpeg', recursive=True)

for filepath in images:
    filename = os.path.basename(filepath)
    print('Carregando o arquivo ' + filename)
    files.append(filename) 

# Gera o relatorio
imgCounter = 0
numImgs = len(files)

while(imgCounter < numImgs):
    addTitle()
    generateTable(imgCounter)
    document.add_page_break()

# Salva documento no PC
document.save('Relatorio.docx')

print("\nDocumento docx gerado.\n")

input("Pressione ENTER para fechar")